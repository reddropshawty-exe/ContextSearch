"""Экстрактор DOCX на базе python-docx."""
from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument

from domain.interfaces import TextExtractor


class DocxExtractor(TextExtractor):
    """Извлекает текст из источников DOCX."""

    def extract(self, source: bytes | str) -> str:
        if isinstance(source, bytes):
            doc = DocxDocument(BytesIO(source))
            return _collect_docx_text(doc)
        path = Path(source)
        if path.exists():
            doc = DocxDocument(path)
            return _collect_docx_text(doc)
        return source


def _collect_docx_text(doc: DocxDocument) -> str:
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


__all__ = ["DocxExtractor"]
